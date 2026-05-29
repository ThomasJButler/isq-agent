"""
DOCX renderer (Plan 9) — typeset compliance report from the canonical envelope.

Consumes the canonical envelope ({questionnaire_meta, answers, summary_metrics}) and
typesets a clean Northstar Labs response document: a title block, a summary table, an
optional run-level banner (all-flagged or all-failed), then one block per question/answer
with flagged answers carrying a [REVIEW] badge + review reason. Style constants live in
app.render.shared.

python-docx note: only body paragraphs (doc.add_paragraph) are visible to readers
scanning doc.paragraphs — text inside tables/headers is not. So the asserted content
(questions, answers, banner) is emitted as body paragraphs; the table carries the
summary visual.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.render.shared import (
    ALL_FAILED_BODY,
    ALL_FAILED_HEADLINE,
    ALL_FLAGGED_BODY,
    ALL_FLAGGED_HEADLINE,
    GREY_CITATION,
    HEADER_SIZE_PT,
    NAVY,
    RED_REVIEW,
    REVIEW_BADGE,
    format_currency,
    format_duration,
    set_cell_border,
    style_run,
)


def render_docx(canonical: dict, output_path: str) -> str:
    """Render the canonical envelope to a DOCX file. Returns the output path."""
    meta = canonical.get("questionnaire_meta", {})
    answers = canonical.get("answers", [])
    summary = canonical.get("summary_metrics", {})

    doc = Document()
    _add_page_branding(doc)
    _add_title(doc, meta)
    _add_summary_table(doc, meta, summary)
    if summary.get("banner") == "all_flagged":
        _add_banner(doc, ALL_FLAGGED_HEADLINE, ALL_FLAGGED_BODY)
    elif summary.get("banner") == "all_failed":
        _add_banner(doc, ALL_FAILED_HEADLINE, ALL_FAILED_BODY)
    _add_answers(doc, answers)

    doc.save(output_path)
    return output_path


def _add_page_branding(doc) -> None:
    """Small navy 'Northstar Labs' mark in the page header (decorative)."""
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("Northstar Labs"), size_pt=8, color=NAVY)


def _add_title(doc, meta: dict) -> None:
    """Title block: response heading, source filename, completion timestamp."""
    origin = meta.get("origin", "the requester")
    title = doc.add_paragraph()
    style_run(
        title.add_run(f"NORTHSTAR LABS — RESPONSE TO {origin}"),
        size_pt=HEADER_SIZE_PT["h1"],
        color=NAVY,
        bold=True,
    )

    if meta.get("filename"):
        style_run(doc.add_paragraph().add_run(meta["filename"]), color=GREY_CITATION)
    if meta.get("completed_at"):
        style_run(
            doc.add_paragraph().add_run(f"Completed: {meta['completed_at']}"),
            color=GREY_CITATION,
        )


def _add_summary_table(doc, meta: dict, summary: dict) -> None:
    """Heading + a bordered 4-row summary table (questions, flagged, cost, time)."""
    heading = doc.add_paragraph()
    style_run(
        heading.add_run("SUMMARY"),
        size_pt=HEADER_SIZE_PT["h2"],
        color=NAVY,
        bold=True,
    )

    rows = [
        ("Total questions:", str(meta.get("total_questions", 0))),
        ("Flagged for review:", str(summary.get("questions_flagged_for_review", 0))),
        ("Total cost:", format_currency(summary.get("total_cost_usd", 0.0))),
        ("Total time:", format_duration(summary.get("total_latency_ms", 0))),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for row_idx, (label, value) in enumerate(rows):
        for col_idx, text in enumerate((label, value)):
            cell = table.cell(row_idx, col_idx)
            set_cell_border(cell)
            style_run(cell.paragraphs[0].add_run(text), bold=col_idx == 0)


def _add_banner(doc, headline_text: str, body_text: str) -> None:
    """Run-level banner as body paragraphs (visible to doc.paragraphs readers)."""
    headline = doc.add_paragraph()
    style_run(
        headline.add_run(headline_text),
        size_pt=HEADER_SIZE_PT["h2"],
        color=RED_REVIEW,
        bold=True,
    )
    style_run(doc.add_paragraph().add_run(body_text), color=GREY_CITATION, italic=True)


def _add_answers(doc, answers: list[dict]) -> None:
    """One block per question/answer; flagged answers get a badge + review reason."""
    heading = doc.add_paragraph()
    style_run(
        heading.add_run("QUESTIONS & ANSWERS"),
        size_pt=HEADER_SIZE_PT["h2"],
        color=NAVY,
        bold=True,
    )

    if not answers:
        style_run(
            doc.add_paragraph().add_run("No questions answered."),
            color=GREY_CITATION,
            italic=True,
        )
        return

    for position, answer in enumerate(answers, start=1):
        _add_answer(doc, position, answer)


def _add_answer(doc, position: int, answer: dict) -> None:
    """Render a single question/answer block."""
    confidence = answer.get("confidence") or {}
    needs_review = confidence.get("needs_review", False)

    question = doc.add_paragraph()
    style_run(
        question.add_run(f"Q{position}: {answer.get('question_text', '')}"),
        size_pt=HEADER_SIZE_PT["h3"],
        color=NAVY,
        bold=True,
    )
    if needs_review:
        style_run(question.add_run(REVIEW_BADGE), color=RED_REVIEW, bold=True)

    style_run(doc.add_paragraph().add_run(answer.get("answer", "")))

    if needs_review and confidence.get("review_reason"):
        style_run(
            doc.add_paragraph().add_run(
                f"Review reason: {confidence['review_reason']}"
            ),
            color=GREY_CITATION,
            italic=True,
        )

    _add_confidence_line(doc, confidence, answer.get("citations", []))


def _add_confidence_line(doc, confidence: dict, citations: list[dict]) -> None:
    """The small grey 'Confidence: X | Citations: ...' line under an answer."""
    score = confidence.get("score")
    score_text = f"{score:.2f}" if isinstance(score, (int, float)) else "n/a"
    source_ids = ", ".join(c.get("source_id", "") for c in citations) or "none"
    style_run(
        doc.add_paragraph().add_run(
            f"Confidence: {score_text} | Citations: {source_ids}"
        ),
        size_pt=9,
        color=GREY_CITATION,
    )
