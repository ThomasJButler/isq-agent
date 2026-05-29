"""
Tests for the DOCX renderer (Plan 9).

Written FIRST per TDD discipline. Implementation in app/render/render_docx.py
follows. The renderer typesets the canonical envelope (questionnaire_meta + answers
+ summary_metrics) into a clean compliance report: a summary table, one block per
question/answer, a [REVIEW] badge + review reason on flagged answers, and an
all-flagged banner.

Note (python-docx): doc.paragraphs returns BODY paragraphs only — text inside
tables/headers/footers is not included. So asserted strings must be emitted as body
paragraphs; the summary table only needs to exist (len(doc.tables) >= 1).
"""

import os

import pytest
from docx import Document

from app.render.render_docx import render_docx


@pytest.fixture
def canonical_one_answer():
    """Minimal canonical envelope with one strong, unflagged answer."""
    return {
        "questionnaire_meta": {
            "origin": "Sunflowers Charity",
            "filename": "Sunflowers_Charity_Supplier_ISQ_Questionnaire.pdf",
            "received_at": "2026-05-25T19:30:00Z",
            "completed_at": "2026-05-25T19:30:42Z",
            "total_questions": 1,
        },
        "answers": [
            {
                "question_id": "sun-q01",
                "question_text": "Do you maintain a formal Information Security Policy?",
                "answer": (
                    "Yes. Northstar Labs maintains a formal Information Security "
                    "Policy which is reviewed annually..."
                ),
                "citations": [
                    {
                        "source_id": "isp-s0",
                        "text_snippet": "Northstar Labs is a UK-based...",
                    },
                ],
                "confidence": {
                    "score": 0.91,
                    "needs_review": False,
                    "review_reason": None,
                },
                "metrics": {
                    "tokens_in": 1240,
                    "tokens_out": 95,
                    "cost_usd": 0.0042,
                    "latency_ms": 1820,
                },
            }
        ],
        "summary_metrics": {
            "total_cost_usd": 0.0042,
            "total_tokens": 1335,
            "total_latency_ms": 1820,
            "questions_flagged_for_review": 0,
            "banner": None,
        },
    }


@pytest.fixture
def canonical_with_flagged_answer():
    """Canonical envelope with one flagged answer and an all_flagged banner."""
    return {
        "questionnaire_meta": {
            "origin": "Blackridge Wind Energy",
            "filename": "Blackridge.pdf",
            "received_at": "2026-05-25T19:30:00Z",
            "completed_at": "2026-05-25T19:30:42Z",
            "total_questions": 1,
        },
        "answers": [
            {
                "question_id": "bla-q01",
                "question_text": "How is privileged access to OT systems controlled?",
                "answer": "Northstar Labs does not operate operational technology systems.",
                "citations": [],
                "confidence": {
                    "score": 0.55,
                    "needs_review": True,
                    "review_reason": "Scope mismatch — Northstar is software-only.",
                },
                "metrics": {
                    "tokens_in": 800,
                    "tokens_out": 60,
                    "cost_usd": 0.003,
                    "latency_ms": 1200,
                },
            }
        ],
        "summary_metrics": {
            "total_cost_usd": 0.003,
            "total_tokens": 860,
            "total_latency_ms": 1200,
            "questions_flagged_for_review": 1,
            "banner": "all_flagged",
        },
    }


def _body_text(path):
    """Join the body paragraph text of a saved DOCX (tables/headers excluded)."""
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _table_text(path):
    """Join all cell text across every table in a saved DOCX."""
    doc = Document(str(path))
    cells = []
    for table in doc.tables:
        for row in table.rows:
            cells.extend(cell.text for cell in row.cells)
    return "\n".join(cells)


class TestRenderDocx:
    def test_writes_file_to_path(self, canonical_one_answer, tmp_path):
        """render_docx writes a file and returns the path it wrote."""
        output = tmp_path / "output.docx"
        result_path = render_docx(canonical_one_answer, str(output))
        assert os.path.exists(result_path)
        assert result_path == str(output)

    def test_output_is_valid_docx(self, canonical_one_answer, tmp_path):
        """The written file opens cleanly as a Word document."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert Document(str(output)) is not None

    def test_includes_question_text(self, canonical_one_answer, tmp_path):
        """The question text appears in the body."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert "Do you maintain a formal Information Security Policy?" in _body_text(
            output
        )

    def test_includes_answer_text(self, canonical_one_answer, tmp_path):
        """The answer text appears in the body."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert (
            "Northstar Labs maintains a formal Information Security Policy"
            in _body_text(output)
        )

    def test_includes_requester_origin(self, canonical_one_answer, tmp_path):
        """The requester origin (questionnaire_meta.origin) appears in the title block."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert "Sunflowers Charity" in _body_text(output)

    def test_includes_summary_table(self, canonical_one_answer, tmp_path):
        """A summary table is present (the visual summary block)."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        doc = Document(str(output))
        assert len(doc.tables) >= 1

    def test_summary_table_shows_question_count(self, canonical_one_answer, tmp_path):
        """The summary table reports the total question count."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert "Total questions" in _table_text(output)

    def test_includes_citation_source_id(self, canonical_one_answer, tmp_path):
        """Citation source_ids are listed under the answer."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        assert "isp-s0" in _body_text(output)

    def test_marks_flagged_answer_with_review_badge(
        self, canonical_with_flagged_answer, tmp_path
    ):
        """A flagged answer carries a REVIEW badge in the body."""
        output = tmp_path / "flagged.docx"
        render_docx(canonical_with_flagged_answer, str(output))
        assert "REVIEW" in _body_text(output)

    def test_includes_review_reason_for_flagged(
        self, canonical_with_flagged_answer, tmp_path
    ):
        """A flagged answer includes its review reason in the body."""
        output = tmp_path / "flagged.docx"
        render_docx(canonical_with_flagged_answer, str(output))
        assert "Scope mismatch" in _body_text(output)

    def test_unflagged_answer_has_no_review_badge(self, canonical_one_answer, tmp_path):
        """An unflagged answer does not get a REVIEW badge."""
        output = tmp_path / "clean.docx"
        render_docx(canonical_one_answer, str(output))
        assert "REVIEW" not in _body_text(output)

    def test_all_flagged_banner_present(self, canonical_with_flagged_answer, tmp_path):
        """The all_flagged banner headline renders as a body paragraph."""
        output = tmp_path / "all_flagged.docx"
        render_docx(canonical_with_flagged_answer, str(output))
        assert "ALL ANSWERS FLAGGED" in _body_text(output).upper()

    def test_all_failed_banner_present(self, canonical_with_flagged_answer, tmp_path):
        """An all_failed run renders its own banner, not just the all_flagged one."""
        canonical = canonical_with_flagged_answer
        canonical["summary_metrics"]["banner"] = "all_failed"
        output = tmp_path / "all_failed.docx"
        render_docx(canonical, str(output))
        assert "ALL ANSWERS FAILED" in _body_text(output).upper()

    def test_no_banner_when_not_all_flagged(self, canonical_one_answer, tmp_path):
        """No banner when summary_metrics.banner is None."""
        output = tmp_path / "no_banner.docx"
        render_docx(canonical_one_answer, str(output))
        assert "ALL ANSWERS FLAGGED" not in _body_text(output).upper()

    def test_no_matrix_terminology_in_output(self, canonical_one_answer, tmp_path):
        """Defence-in-depth: the renderer never propagates Matrix terms."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        full_text = _body_text(output).lower()
        forbidden = [
            "morpheus",
            "the matrix",
            "neo ",
            "white rabbit",
            "redpill",
            "bluepill",
        ]
        for term in forbidden:
            assert term not in full_text, (
                f"Matrix term '{term}' leaked into DOCX output"
            )

    def test_uses_calibri_body_font(self, canonical_one_answer, tmp_path):
        """The first body run uses Calibri (compliance palette body font)."""
        output = tmp_path / "output.docx"
        render_docx(canonical_one_answer, str(output))
        doc = Document(str(output))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.name:
                    assert run.font.name == "Calibri", (
                        f"Unexpected font: {run.font.name}"
                    )
                    break
            break

    def test_handles_empty_answers(self, tmp_path):
        """An empty run still renders a valid file (summary only)."""
        empty_canonical = {
            "questionnaire_meta": {
                "origin": "Test",
                "filename": "test.pdf",
                "received_at": "",
                "completed_at": "",
                "total_questions": 0,
            },
            "answers": [],
            "summary_metrics": {
                "total_cost_usd": 0.0,
                "total_tokens": 0,
                "total_latency_ms": 0,
                "questions_flagged_for_review": 0,
                "banner": None,
            },
        }
        output = tmp_path / "empty.docx"
        result = render_docx(empty_canonical, str(output))
        assert os.path.exists(result)
