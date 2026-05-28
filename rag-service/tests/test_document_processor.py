"""
Tests for the document processor module (Plan 4).
Written FIRST per TDD discipline. Implementation in app/utils/document_processor.py follows.
"""

import pytest


# Test fixtures


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Create a minimal valid PDF for testing."""
    pdf_path = tmp_path / "test_policy.pdf"
    # Minimal PDF structure (PDF 1.4 header + simple content)
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
50 750 Td
(Test PDF content) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000317 00000 n
trailer
<< /Size 5 /Root 1 0 R >>
startxref
410
%%EOF
"""
    pdf_path.write_bytes(pdf_content)
    return pdf_path


@pytest.fixture
def corrupted_pdf_path(tmp_path):
    """Create a malformed PDF for error handling tests."""
    pdf_path = tmp_path / "corrupted.pdf"
    pdf_path.write_bytes(b"Not a real PDF file")
    return pdf_path


@pytest.fixture
def sample_docx_path(tmp_path):
    """Create a minimal DOCX for testing."""
    # Note: This requires python-docx to be installed
    # For now, we'll create a placeholder — implementation will use python-docx
    docx_path = tmp_path / "test_policy.docx"

    # Create a minimal valid DOCX using python-docx
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("This is test DOCX content.")
        doc.add_paragraph("Second paragraph in DOCX.")
        doc.save(str(docx_path))
    except ImportError:
        pytest.skip("python-docx not installed")

    return docx_path


@pytest.fixture
def sample_xlsx_path(tmp_path):
    """Create a minimal XLSX for testing historical ISQ parsing."""
    xlsx_path = tmp_path / "test_isq.xlsx"

    # Create a minimal valid XLSX using openpyxl
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active

        # Header row
        ws["A1"] = "Question"
        ws["B1"] = "Answer"

        # Data rows
        ws["A2"] = "Do you use MFA?"
        ws["B2"] = "Yes, MFA is mandatory for all cloud platforms."

        ws["A3"] = "Where is customer data stored?"
        ws["B3"] = "Customer data is stored in UK data centers."

        wb.save(str(xlsx_path))
    except ImportError:
        pytest.skip("openpyxl not installed")

    return xlsx_path


@pytest.fixture
def unsupported_file_path(tmp_path):
    """Create an unsupported file type for rejection tests."""
    txt_path = tmp_path / "test.txt"
    txt_path.write_text("Plain text file content")
    return txt_path


# Tests


def test_process_pdf_extracts_text(sample_pdf_path):
    """PDF file → returns dict with text, page_count, pages[]."""
    from app.utils.document_processor import process_document

    result = process_document(sample_pdf_path)

    assert isinstance(result, dict)
    assert "text" in result
    assert "page_count" in result
    assert "pages" in result
    assert isinstance(result["text"], str)
    assert len(result["text"]) > 0
    assert result["page_count"] >= 1
    assert isinstance(result["pages"], list)
    assert len(result["pages"]) == result["page_count"]


def test_process_docx_extracts_text(sample_docx_path):
    """DOCX file → returns dict with text."""
    from app.utils.document_processor import process_document

    result = process_document(sample_docx_path)

    assert isinstance(result, dict)
    assert "text" in result
    assert isinstance(result["text"], str)
    assert len(result["text"]) > 0
    assert "DOCX content" in result["text"]


def test_process_xlsx_extracts_rows(sample_xlsx_path):
    """XLSX file → returns dict with rows[], each row having question_text if column matches."""
    from app.utils.document_processor import process_document

    result = process_document(sample_xlsx_path)

    assert isinstance(result, dict)
    assert "rows" in result
    assert isinstance(result["rows"], list)
    assert len(result["rows"]) >= 2  # Header + at least 1 data row

    # Check that rows contain question/answer data
    # (Implementation should detect Question/Answer columns)
    data_rows = [r for r in result["rows"] if r.get("question_text")]
    assert len(data_rows) > 0

    first_row = data_rows[0]
    assert "question_text" in first_row
    assert "answer_text" in first_row


def test_process_rejects_unsupported_type(unsupported_file_path):
    """.txt or other unsupported → raises ValueError."""
    from app.utils.document_processor import process_document

    with pytest.raises(ValueError, match="Unsupported file type"):
        process_document(unsupported_file_path)


def test_process_handles_corrupted_pdf(corrupted_pdf_path):
    """Malformed PDF → raises DocumentProcessingError with clear message."""
    from app.utils.document_processor import process_document, DocumentProcessingError

    with pytest.raises(DocumentProcessingError, match="Failed to process PDF"):
        process_document(corrupted_pdf_path)


def test_process_preserves_page_numbers(sample_pdf_path):
    """Multi-page PDF → each page recorded with page number."""
    from app.utils.document_processor import process_document

    result = process_document(sample_pdf_path)

    assert "pages" in result
    for page_data in result["pages"]:
        assert "page_number" in page_data
        assert "text" in page_data
        assert isinstance(page_data["page_number"], int)
        assert page_data["page_number"] >= 1


def test_process_pdf_handles_pages_with_no_extractable_text(tmp_path):
    """pypdf returns None for image-only/scanned/blank pages → no crash, empty text."""
    from unittest.mock import MagicMock, patch

    from app.utils.document_processor import process_document

    pdf_path = tmp_path / "image_only.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 placeholder")

    # Page 1 is image-only (extract_text() -> None), page 2 has real text.
    page_none = MagicMock()
    page_none.extract_text.return_value = None
    page_text = MagicMock()
    page_text.extract_text.return_value = "Real policy text"
    fake_reader = MagicMock()
    fake_reader.pages = [page_none, page_text]

    with patch("app.utils.document_processor.PdfReader", return_value=fake_reader):
        result = process_document(pdf_path)

    assert result["page_count"] == 2
    assert result["pages"][0]["text"] == ""  # None coerced to empty string
    assert "Real policy text" in result["text"]
